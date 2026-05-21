"""Extract references and in-text citations from various input types.

We accept four input formats:
  * .pdf       -- parsed with pdfplumber; references found by section heading
  * .docx      -- parsed with python-docx; references found by section heading
  * .bib       -- BibTeX, parsed with bibtexparser
  * .ris       -- RIS, parsed with rispy
  * .txt/.md   -- plain text, each non-empty paragraph is one citation

The extractor is intentionally lenient: it returns ParsedReference records that
verify.py can later confirm or correct against authoritative APIs. Garbage in
is still extracted, then flagged downstream.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple, Optional

from .models import ParsedReference, Author, CitationType
from .normalize import (
    extract_doi, extract_arxiv_id, extract_pmid, extract_year,
    is_institutional_author,
)


# ---------- in-text citation extraction (works on the body of the paper) ----

# Surname characters: A-Z plus extended Latin (Müller, García, Łukasiewicz).
_SN_CHARS = r"A-Za-z\-'À-ÖØ-öø-ÿĀ-žŁł"
_SN_CAP = r"[A-ZÀ-ÖØ-ÞĀĂ-ŻŁ]"

# Institutional-author body in parenthetical citation: multiple capitalized
# words connected by short connectives. E.g. 'Bureau of Justice Statistics',
# 'Centers for Disease Control and Prevention', 'U.S. Department of Justice'.
_INST_CAP_WORD = rf"{_SN_CAP}\.?[{_SN_CHARS}\.]{{0,40}}"   # 'Bureau', 'U.S.', 'CDC'
_INST_CONN = r"(?:of|for|and|the|on|in|to|\&)"
_INSTITUTIONAL_BODY = rf"{_INST_CAP_WORD}(?:\s+(?:{_INST_CONN}\s+)?{_INST_CAP_WORD}){{1,8}}"

_INTEXT_PATTERNS = [
    # Parenthetical personal author-year:
    #   (Smith, 2017)  (Smith et al., 2017)  (Smith and Jones, 2017)  (Smith & Jones, 2017)
    re.compile(
        r"\("
        rf"{_SN_CAP}[{_SN_CHARS}]{{1,40}}"
        rf"(?:\s+(?:and|&)\s+{_SN_CAP}[{_SN_CHARS}]{{1,40}})?"
        r"(?:\s+et\s+al\.?)?"
        r"(?:,\s*|\s+)"
        r"\d{4}[a-z]?"
        r"\)"
    ),
    # Parenthetical institutional author-year:
    #   (Bureau of Justice Statistics, 2024)
    #   (Centers for Disease Control and Prevention, 2024)
    re.compile(
        r"\("
        rf"{_INSTITUTIONAL_BODY}"
        r",\s*\d{4}[a-z]?"
        r"\)"
    ),
    # Numeric: [1] [12] [1, 2, 3] [1-3]
    re.compile(r"\[(\d{1,3}(?:\s*[,\-–]\s*\d{1,3})*)\]"),
    # Narrative personal author-year:
    #   Smith (2017)            Smith et al. (2017)
    #   Smith and Jones (2017)  Smith & Jones (2017)
    re.compile(
        rf"\b{_SN_CAP}[{_SN_CHARS}]{{1,40}}"
        rf"(?:\s+(?:and|&)\s+{_SN_CAP}[{_SN_CHARS}]{{1,40}})?"
        r"(?:\s+et\s+al\.?)?"
        r"\s*\(\d{4}[a-z]?\)"
    ),
    # Narrative institutional author-year:
    #   Centers for Disease Control (2024)
    re.compile(
        rf"\b{_INSTITUTIONAL_BODY}\s*\(\d{{4}}[a-z]?\)"
    ),
]


# A whole parenthetical that COULD contain multiple semicolon-separated
# citations: (Smith, 2017; Jones, 2018; Bureau of Justice Statistics, 2024).
# Matched first so we can split it; individual citation patterns then run
# over the segments AND the original text for narrative forms.
_MULTI_CITE_PARENS = re.compile(
    r"\("
    r"(?:[^()]{1,500}?)"
    r"\)"
)


def _is_multi_cite(content: str) -> bool:
    """Heuristic: does the content of a parenthetical look like multiple
    semicolon-separated author-year citations?"""
    if ";" not in content:
        return False
    parts = [p.strip() for p in content.split(";") if p.strip()]
    if len(parts) < 2:
        return False
    # At least 2 segments must each contain a 4-digit year. Otherwise it's
    # a single citation with an internal semicolon (e.g., 'a; b in series').
    year_segments = sum(1 for p in parts if re.search(r"\b(?:19|20)\d{2}\b", p))
    return year_segments >= 2


def extract_intext_citations(text: str) -> List[str]:
    """Return de-duplicated raw in-text citation strings.

    Splits multi-citation parentheticals like '(Smith, 2017; Jones, 2018)'
    into individual citations.
    """
    seen: List[str] = []

    def _add(s: str) -> None:
        s = s.strip()
        if s and s not in seen:
            seen.append(s)

    # First pass: extract whole parentheticals that look like multi-cite
    # and split on semicolons. Each segment is wrapped back in parens so it
    # matches the parenthetical patterns later.
    used_spans: List[Tuple[int, int]] = []
    for m in _MULTI_CITE_PARENS.finditer(text):
        inner = m.group(0)[1:-1]  # strip outer parens
        if not _is_multi_cite(inner):
            continue
        used_spans.append(m.span())
        for seg in inner.split(";"):
            seg = seg.strip()
            if seg:
                _add(f"({seg})")

    # Second pass: run the standard patterns on the rest of the text.
    for pat in _INTEXT_PATTERNS:
        for m in pat.finditer(text):
            # Skip matches that fall inside a parenthetical we already split.
            if any(s <= m.start() < e for s, e in used_spans):
                continue
            _add(m.group(0))
    return seen


# ---------- reference list location -----------------------------------------

_REF_HEADING = re.compile(
    r"^\s*(references|bibliography|works\s+cited|literature\s+cited|citations|reference\s+list)\s*:?\s*$",
    re.IGNORECASE,
)
_NEXT_SECTION_HEADING = re.compile(
    r"^\s*(appendix|appendices|supplementary|supporting\s+information|acknowledg(e)?ments?|"
    r"figures?|tables?|index|notes?|glossary|abbreviations?)\b",
    re.IGNORECASE,
)

# Appendix subheadings: 'A.', 'B.', 'C.', 'A. Title', 'B Background', etc.
# These commonly appear after the References section in academic outlines.
_APPENDIX_LETTER_HEADING = re.compile(
    r"^\s*[A-Z]\.\s+[A-Z]"
)


def _find_references_block(lines: List[str]) -> List[str]:
    """Given a list of body lines, return the lines inside the references section.

    If no explicit heading is found, fall back to the whole input — the user
    may have pasted a bare bibliography.
    """
    start = None
    for i, line in enumerate(lines):
        if _REF_HEADING.match(line.strip()):
            start = i + 1
            break
    if start is None:
        return lines
    end = len(lines)
    for j in range(start, len(lines)):
        line = lines[j].strip()
        if _NEXT_SECTION_HEADING.match(line):
            end = j
            break
        # Appendix letter heading ('C. Responsible AI ...') -- but only if it
        # appears AFTER at least one plausible reference (avoid clipping the
        # whole list when the very first ref happens to look like 'A.').
        if j > start and _APPENDIX_LETTER_HEADING.match(line):
            end = j
            break
    return lines[start:end]


# ---------- split a references block into individual entries ---------------

_NUMBERED_PREFIX = re.compile(r"^\s*(?:\[(\d{1,3})\]|(\d{1,3})[.)])\s+")
_AUTHOR_LEAD = re.compile(r"^[A-Z][A-Za-z\-']+,?\s+[A-Z]\.")  # 'Smith, J.' or 'Smith J.'


def split_into_entries(block_lines: List[str]) -> List[str]:
    """Group raw lines into one string per reference entry.

    Strategy, in order:
      1. Numbered prefix ([1] / 1. / 1)) -- split on the numbers.
      2. Treat each non-empty line as a candidate entry if MOST lines look
         like standalone references (start with a capital letter or digit AND
         contain a 4-digit year). This handles docx/PDF inputs where each
         reference is its own paragraph but without blank lines between them.
      3. Blank-line paragraph splitting (the classic plain-text case).
      4. Author-lead heuristic for one big blob.
    """
    text = "\n".join(block_lines)
    if not text.strip():
        return []

    # Strategy 1: numbered
    numbered = list(_NUMBERED_PREFIX.finditer(text))
    if len(numbered) >= 2:
        entries = []
        for i, m in enumerate(numbered):
            start = m.end()
            end = numbered[i + 1].start() if i + 1 < len(numbered) else len(text)
            entry = text[start:end].strip()
            entry = re.sub(r"\s+", " ", entry)
            if entry:
                entries.append(entry)
        return entries

    # Strategy 2: one paragraph (line) per reference. Detect this by checking
    # what fraction of non-empty lines independently look like a citation.
    non_empty = [l.strip() for l in block_lines if l.strip()]
    if len(non_empty) >= 2:
        looks_like_ref = lambda s: (
            bool(re.match(r"^[A-Z0-9\(\[]", s))
            and bool(re.search(r"\(\d{4}[a-z]?\)|\b(19|20)\d{2}\b", s))
            and len(s) > 30
        )
        n_ref_like = sum(1 for l in non_empty if looks_like_ref(l))
        if n_ref_like >= max(2, int(0.6 * len(non_empty))):
            # Most lines look like references on their own. Take them as-is,
            # but fold any continuation line (one that doesn't itself look like
            # a reference) onto the previous entry.
            entries: List[str] = []
            for line in non_empty:
                if looks_like_ref(line) or not entries:
                    entries.append(line)
                else:
                    entries[-1] = entries[-1] + " " + line
            entries = [re.sub(r"\s+", " ", e).strip() for e in entries]
            return entries

    # Strategy 3: blank-line paragraph splitting.
    raw_paras = re.split(r"\n\s*\n", text)
    entries = []
    for para in raw_paras:
        para = re.sub(r"\s+", " ", para).strip()
        if para:
            entries.append(para)
    if len(entries) > 1:
        return entries

    # Strategy 4 fallback: split on author-lead boundaries within one big blob.
    parts = re.split(r"(?<=[\.\]])\s+(?=[A-Z][A-Za-z\-']+,\s+[A-Z]\.)", text.replace("\n", " "))
    return [re.sub(r"\s+", " ", p).strip() for p in parts if p.strip()]


# ---------- per-entry field guessing ----------------------------------------

_QUOTED_TITLE = re.compile(r"[\"“]([^\"”]{4,300})[\"”]")
_BARE_TITLE_AFTER_YEAR = re.compile(r"\(\d{4}[a-z]?\)\.\s*([^.]{4,300})\.")
_BARE_TITLE_AFTER_AUTHORS = re.compile(r"\.\s+([A-Z][^.]{4,300})\.")

# Author block: everything up to the first parenthesized 4-digit year or the
# APA 'no date' marker '(n.d.)'. We require the '(YYYY' shape (open paren,
# then digits) so bare 4-digit numbers inside a title don't truncate the list.
# Don't require the closing paren -- APA allows extended dates like '(2026,
# March 11)' or '(2024a)'.
_AUTHOR_BLOCK = re.compile(
    r"^(.+?)(?=\s*\(\d{4}\b|\s*\(n\.?d\.?\)|[\"“])",
    re.DOTALL,
)


def _parse_authors(s: str) -> List[Author]:
    """Parse an author block, handling both personal AND institutional authors.

    For institutional authors ('Bureau of Justice Statistics.'), return a
    single Author whose `family` is the full institution name and `given`
    is empty. This avoids the bug where 'Bureau of Justice Statistics' was
    parsed as family='Statistics', given='Bureau of Justice'.
    """
    s = s.strip().rstrip(",.;")
    if not s:
        return []
    # Strip 'et al.' / 'and others' / 'and colleagues' before parsing so they
    # don't end up as authors named 'al'.
    s = re.sub(r"[,\s]+et\s+al\.?\b\s*$", "", s, flags=re.IGNORECASE)
    s = re.sub(r"[,\s]+and\s+(?:others|colleagues)\b\s*$", "", s, flags=re.IGNORECASE)
    s = s.strip().rstrip(",.;")
    if not s:
        return []

    # Institutional author? Single entity, return it as one Author.
    if is_institutional_author(s):
        return [Author(family=s, given="")]

    # Some references have a mix: 'Smith J. & World Health Organization' --
    # detect by splitting on 'and'/'&' and checking each side independently.
    # Only do this split at the top level (not inside an institutional name).
    if re.search(r"\s+(and|&)\s+", s):
        # Try to split into person+institution candidates.
        major_split = re.split(r"\s+(?:and|&)\s+", s)
        all_institutional = all(is_institutional_author(p.strip()) for p in major_split if p.strip())
        any_institutional = any(is_institutional_author(p.strip()) for p in major_split if p.strip())
        if all_institutional:
            return [Author(family=p.strip().rstrip(",.;"), given="") for p in major_split if p.strip()]
        if any_institutional and len(major_split) <= 3:
            out: List[Author] = []
            for p in major_split:
                p = p.strip().rstrip(",.;")
                if not p:
                    continue
                if is_institutional_author(p):
                    out.append(Author(family=p, given=""))
                else:
                    out.extend(_parse_authors(p))  # recurse for the person-author side
            return out

    # Personal authors: 'Smith, J., Jones, K., & Brown, L.' style.
    s2 = re.sub(r"\s+(and|&)\s+", ", ", s)
    parts = [p.strip() for p in s2.split(",") if p.strip()]
    fused: List[str] = []
    i = 0
    while i < len(parts):
        cur = parts[i]
        if i + 1 < len(parts) and re.fullmatch(r"(?:[A-Z]\.?\s*){1,4}", parts[i + 1]):
            fused.append(f"{cur}, {parts[i + 1]}")
            i += 2
        else:
            fused.append(cur)
            i += 1
    return [Author.from_string(p) for p in fused[:50]]


def _parse_one_entry(raw: str) -> ParsedReference:
    ref = ParsedReference(raw_text=raw)

    ref.doi = extract_doi(raw)
    ref.arxiv_id = extract_arxiv_id(raw)
    ref.pmid = extract_pmid(raw)
    ref.year = extract_year(raw)

    # URL (best-effort, only if no DOI was found inline)
    m_url = re.search(r"https?://\S+", raw)
    if m_url and not ref.doi:
        ref.url = m_url.group(0).rstrip(".,);")

    # Title -- try quoted, then post-year, then post-authors heuristic. Each
    # candidate is sanity-checked: short chunks ending in 'comma + single
    # capital letter' are author-list fragments, not titles. Chunks that
    # consist only of a URL are also rejected (regex can otherwise capture
    # 'https://doi' as the title when year-parens precede a DOI URL).
    def _looks_like_title(s: str) -> bool:
        s = s.strip().rstrip(".,;")
        # Floor of 6 admits legitimately short titles ('Deep learning',
        # 'Attention') while rejecting stray fragments.
        if len(s) < 6:
            return False
        # Author chunk: 'Chatterjee, S' / 'Brady, P' / 'Smith, J. K'
        if re.fullmatch(r"[A-Za-z\-']+(?:,\s*[A-Z]\.?)+", s):
            return False
        # URL fragment
        if s.lower().startswith(("http://", "https://", "www.")) or "://" in s:
            return False
        # Mostly punctuation / initials (e.g. 'A. B., C. D., E. F'): >50% of
        # tokens are 1-2 letter capitals.
        tokens = re.findall(r"\S+", s)
        if tokens:
            short_caps = sum(1 for t in tokens if re.fullmatch(r"[A-Z]{1,2}\.?", t.rstrip(",.;")))
            if short_caps / len(tokens) > 0.5:
                return False
        return True

    for pat in (_QUOTED_TITLE, _BARE_TITLE_AFTER_YEAR, _BARE_TITLE_AFTER_AUTHORS):
        for m in pat.finditer(raw):
            cand = m.group(1).strip().rstrip(".,;")
            if _looks_like_title(cand):
                ref.title = cand
                break
        if ref.title:
            break

    # Authors. We have two candidate strategies:
    #   A. APA-style: everything up to '(YYYY' or '(n.d.)' marker. Works when
    #      the year-in-parens is the author/title boundary.
    #   B. Title-position: everything before the detected title. Works for
    #      AMA-style references where the year appears at the end.
    #
    # We try both and pick whichever produces a sensible result. A 'sensible'
    # author block is shorter than 200 characters, does NOT contain the
    # detected title as a substring, and looks structurally plausible.
    # This handles documents where APA-style '(year)' appears in the middle
    # of an entry (e.g. an AMA citation that includes a parenthesized year
    # somewhere after the title).
    candidates: List[str] = []
    author_block_match = _AUTHOR_BLOCK.match(raw)
    if author_block_match:
        candidates.append(author_block_match.group(1).strip().rstrip(",.;"))
    if ref.title:
        title_pos = raw.find(ref.title)
        if title_pos > 0:
            candidates.append(raw[:title_pos].strip().rstrip(",.;"))

    def _looks_like_author_block(s: str) -> bool:
        if not s or len(s) > 250:
            return False
        if ref.title and ref.title.lower() in s.lower():
            return False
        return True

    for cand in candidates:
        if _looks_like_author_block(cand):
            ref.authors = _parse_authors(cand)
            break

    # Container title heuristic. After the title, the journal name runs up to
    # the volume number. Handles abbreviated AMA/Vancouver journal names that
    # contain internal periods ('J. Am. Med. Inform. Assoc.') and journals
    # that start lowercase ('npj Digit. Med.', 'eLife', 'medRxiv').
    if ref.title:
        after_title = raw.split(ref.title, 1)[-1]
        # Strip stray leading punctuation (malformed close-quotes, periods).
        after_title = after_title.lstrip(' ".”“')
        # Primary: capture words (allowing periods/&) up to the volume number.
        m_container = re.search(
            r"^[\s.\"'“”]*([A-Za-z][A-Za-z.&'\-\s]+?)[,.]?\s+\d{1,4}\s*[,(:;]",
            after_title,
        )
        if m_container:
            ref.container_title = m_container.group(1).strip().rstrip(".")
        else:
            # Fallback: first capitalized chunk before a comma/period.
            m2 = re.search(r"[\.\?!]\s+([A-Z][^,.]{2,120})[,.]", after_title)
            if m2:
                ref.container_title = m2.group(1).strip()

    # Volume / issue / pages. AMA style writes 'Year;Vol(Issue):Pages.' so try
    # both 'Vol(Issue): Pages' (APA) and 'Year;Vol(Issue):Pages' (AMA).
    m_vip = re.search(r"\b(\d{1,4})\s*\((\d{1,4})\)\s*[,:]\s*([0-9eA-Z]+\s*[-–]\s*[0-9eA-Z]+)", raw)
    if m_vip:
        ref.volume, ref.issue, ref.pages = m_vip.group(1), m_vip.group(2), m_vip.group(3).replace(" ", "")
    else:
        m_v = re.search(r"\bvol(?:ume)?\.?\s*(\d{1,4})", raw, re.IGNORECASE)
        if m_v:
            ref.volume = m_v.group(1)
        m_p = re.search(r"\bpp?\.?\s*([0-9eA-Z]+\s*[-–]\s*[0-9eA-Z]+)", raw)
        if m_p:
            ref.pages = m_p.group(1).replace(" ", "")

    # Publisher: AMA grey literature writes 'City, State: Publisher; Year.' or
    # 'City: Publisher; Year.'. Detect this independently of container_title
    # (the container heuristic can mis-fire on city names and we'd lose the
    # publisher signal otherwise).
    m_pub = re.search(
        r"\b[A-Z][A-Za-z]+(?:,\s*[A-Z]{2,})?:\s+([A-Z][^;]{2,80}?);\s*\d{4}",
        raw,
    )
    if m_pub:
        ref.publisher = m_pub.group(1).strip().rstrip(".")
        # If container_title also got captured but it's actually the city
        # name preceding our publisher, clear it -- this is grey lit, not a
        # journal article.
        if ref.container_title and ref.container_title in m_pub.group(0):
            ref.container_title = ""

    # Type guess
    if ref.arxiv_id and not ref.doi:
        ref.type = CitationType.PREPRINT
    elif re.search(r"\bproc(?:eedings)?\b|\bconf(?:erence)?\b", raw, re.IGNORECASE):
        ref.type = CitationType.CONFERENCE_PAPER
    elif re.search(r"\bpress\b|\bpublisher\b|\bbook\b", raw, re.IGNORECASE):
        ref.type = CitationType.BOOK
    elif ref.container_title or ref.doi or ref.pmid or ref.volume:
        ref.type = CitationType.JOURNAL_ARTICLE
    else:
        ref.type = CitationType.UNKNOWN

    return ref


# ---------- public entry points --------------------------------------------

def _body_text(lines: List[str]) -> str:
    """Return the body of the paper -- everything BEFORE the References heading.

    If there's no References heading, returns empty so we don't treat a bare
    bibliography as a paper with in-text citations.
    """
    for i, line in enumerate(lines):
        if _REF_HEADING.match(line.strip()):
            return "\n".join(lines[:i])
    return ""


def _safe_parse_one_entry(raw: str) -> ParsedReference:
    """Crash-proof wrapper around _parse_one_entry. On ANY failure, return a
    minimal record carrying the raw text so it's still surfaced for review
    rather than aborting the whole audit."""
    try:
        return _parse_one_entry(raw)
    except Exception as e:  # pragma: no cover - defensive
        import logging
        logging.getLogger("biblio-check.extract").debug(
            "entry parse failed (%s); keeping raw text only", e
        )
        return ParsedReference(raw_text=raw)


def extract_from_text(text: str, parser_backend: str = "auto") -> Tuple[List[ParsedReference], List[str]]:
    """Extract references and in-text citations from a plain-text blob.

    In-text citations are pulled from the body section only (above the
    References heading), so we don't pollute the cross-check by treating the
    bibliography entries themselves as in-text citations.

    `parser_backend`: 'auto' (use AnyStyle/GROBID if available, else regex),
    'regex' (force the built-in parser), 'anystyle', or 'grobid'.
    """
    lines = text.splitlines()
    block = _find_references_block(lines)
    entries = split_into_entries(block)

    # Route per-entry field parsing through the optional backend abstraction.
    # The backend is auto-detected and falls back to the built-in regex parser
    # if nothing better is installed -- so this NEVER changes whether the tool
    # runs, only how accurately it parses.
    try:
        from .parsers import parse_entries
        refs = parse_entries(entries, backend=parser_backend,
                             regex_parser=_safe_parse_one_entry)
    except Exception as e:  # pragma: no cover - defensive
        import logging
        logging.getLogger("biblio-check.extract").debug(
            "backend parse failed (%s); using regex parser", e
        )
        refs = [_safe_parse_one_entry(e) for e in entries]

    body = _body_text(lines)
    try:
        intext = extract_intext_citations(body) if body else []
    except Exception:  # pragma: no cover - defensive
        intext = []
    return refs, intext


def _extract_with_grobid(path: Path, grobid_url: str) -> Optional[Tuple[List[ParsedReference], List[str]]]:
    """Optional GROBID-backed PDF reference extraction.

    GROBID is a Java service that does structured PDF parsing far better than
    regex-based extraction. It's opt-in via the GROBID_URL env var because it
    requires a running server. If unreachable or returns no useful data, we
    return None and the caller falls back to pdfplumber.

    Reference: https://grobid.readthedocs.io/en/latest/Grobid-service/
    """
    import urllib.request
    import urllib.error
    import xml.etree.ElementTree as ET

    endpoint = grobid_url.rstrip("/") + "/api/processReferences"
    boundary = "----biblio-check-grobid-boundary"
    with open(path, "rb") as fh:
        pdf_bytes = fh.read()
    body_parts = [
        f"--{boundary}\r\n".encode(),
        b'Content-Disposition: form-data; name="input"; filename="' + path.name.encode() + b'"\r\n',
        b"Content-Type: application/pdf\r\n\r\n",
        pdf_bytes,
        f"\r\n--{boundary}\r\n".encode(),
        b'Content-Disposition: form-data; name="consolidateCitations"\r\n\r\n1\r\n',
        f"--{boundary}--\r\n".encode(),
    ]
    body = b"".join(body_parts)
    req = urllib.request.Request(
        endpoint, data=body, method="POST",
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}",
                 "Accept": "application/xml"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60.0) as r:
            xml_text = r.read().decode("utf-8", errors="replace")
    except (urllib.error.HTTPError, urllib.error.URLError, TimeoutError) as e:
        import logging
        logging.getLogger("biblio-check.extract").debug("GROBID unreachable: %s", e)
        return None

    try:
        root = ET.fromstring(xml_text)
    except ET.ParseError:
        return None

    ns = {"tei": "http://www.tei-c.org/ns/1.0"}
    refs: List[ParsedReference] = []
    for bibl in root.iter("{http://www.tei-c.org/ns/1.0}biblStruct"):
        ref = ParsedReference()
        # Title
        title_el = bibl.find(".//tei:title[@level='a']", ns) or bibl.find(".//tei:title", ns)
        if title_el is not None and title_el.text:
            ref.title = title_el.text.strip()
        # Authors
        for pers in bibl.findall(".//tei:author/tei:persName", ns):
            surname = pers.find("tei:surname", ns)
            given = pers.find("tei:forename", ns)
            if surname is not None and surname.text:
                ref.authors.append(Author(
                    family=surname.text.strip(),
                    given=(given.text.strip() if given is not None and given.text else ""),
                ))
        # Year
        date_el = bibl.find(".//tei:date[@type='published']", ns)
        if date_el is not None:
            when = date_el.get("when") or (date_el.text or "")
            m_year = re.match(r"(\d{4})", when)
            if m_year:
                ref.year = int(m_year.group(1))
        # Journal / venue
        j_el = bibl.find(".//tei:title[@level='j']", ns) or bibl.find(".//tei:title[@level='m']", ns)
        if j_el is not None and j_el.text:
            ref.container_title = j_el.text.strip()
        # Volume / issue / pages
        bs = bibl.find(".//tei:biblScope[@unit='volume']", ns)
        if bs is not None and bs.text:
            ref.volume = bs.text.strip()
        bs = bibl.find(".//tei:biblScope[@unit='issue']", ns)
        if bs is not None and bs.text:
            ref.issue = bs.text.strip()
        bs = bibl.find(".//tei:biblScope[@unit='page']", ns)
        if bs is not None:
            if bs.text:
                ref.pages = bs.text.strip()
            elif bs.get("from") and bs.get("to"):
                ref.pages = f"{bs.get('from')}-{bs.get('to')}"
        # DOI
        doi_el = bibl.find(".//tei:idno[@type='DOI']", ns)
        if doi_el is not None and doi_el.text:
            ref.doi = doi_el.text.strip().lower()
        # URL
        ptr = bibl.find(".//tei:ptr", ns)
        if ptr is not None and ptr.get("target"):
            ref.url = ptr.get("target")
        # raw_text fallback for downstream display
        ref.raw_text = ET.tostring(bibl, encoding="unicode")
        if ref.title or ref.authors:
            refs.append(ref)
    if not refs:
        return None
    return refs, []


def extract_from_pdf(path: Path, parser_backend: str = "auto") -> Tuple[List[ParsedReference], List[str]]:
    """Extract from a PDF.

    If the GROBID_URL env var points to a running GROBID server, try that
    first (it does proper structured parsing). On any failure or absence,
    fall back to pdfplumber + the configured text parser backend.
    """
    import os
    grobid_url = os.environ.get("GROBID_URL", "").strip()
    if grobid_url and parser_backend in ("auto", "grobid"):
        try:
            grobid_result = _extract_with_grobid(path, grobid_url)
            if grobid_result:
                return grobid_result
        except Exception:
            pass  # fall through to pdfplumber

    try:
        import pdfplumber
    except ImportError as e:
        raise RuntimeError(
            "pdfplumber is required to read PDF files. Install with: pip install pdfplumber"
        ) from e
    try:
        pages_text: List[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text() or ""
                pages_text.append(txt)
    except Exception as e:
        raise RuntimeError(
            f"Could not read PDF {path.name}. The file may be corrupted, scanned (image-only), "
            f"or password-protected. Try copying the bibliography into a .txt file. "
            f"Underlying error: {e}"
        ) from e
    full_text = "\n".join(pages_text)
    return extract_from_text(full_text, parser_backend=parser_backend)


def _extract_superscript_citations(d) -> List[str]:
    """Walk every run in every paragraph; collect digit-sequences in
    superscript runs as Vancouver/AMA-style numeric citations.

    python-docx flattens superscript runs into inline text when you read
    paragraph.text, so '... environments.{sup}16{/sup}' becomes
    '... environments.16'. We can't distinguish citations from inline numbers
    that way. Reading runs directly preserves the superscript flag.

    Returns a list of citation strings in '[N]' form so they match the same
    downstream parser as bracketed numeric citations.
    """
    out: List[str] = []
    seen: set = set()
    for para in d.paragraphs:
        for run in para.runs:
            try:
                if not run.font.superscript:
                    continue
            except AttributeError:
                continue
            txt = (run.text or "").strip()
            if not txt:
                continue
            # Collect digit-comma-hyphen sequences inside the run text.
            for m in re.finditer(r"\d{1,3}(?:[\-–,]\s*\d{1,3})*", txt):
                key = m.group(0).replace(" ", "")
                citation = f"[{key}]"
                if citation not in seen:
                    seen.add(citation)
                    out.append(citation)
    return out


def extract_from_docx(path: Path, parser_backend: str = "auto") -> Tuple[List[ParsedReference], List[str]]:
    """Extract from a .docx using python-docx.

    In addition to plain-text in-text citation extraction (parenthetical
    author-year and bracketed numeric), this also walks runs to detect
    superscript citations (AMA / Vancouver style), which python-docx's
    paragraph.text flattens into bare inline digits.
    """
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required to read .docx files. Install with: pip install python-docx"
        ) from e
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs]
        sup_citations = _extract_superscript_citations(d)
    except Exception as e:
        raise RuntimeError(
            f"Could not read .docx {path.name}: {e}"
        ) from e
    full_text = "\n".join(paras)
    refs, intext = extract_from_text(full_text, parser_backend=parser_backend)
    # Append superscript citations we couldn't see in the flat text. Skip ones
    # that are already represented by the standard bracketed-numeric pattern.
    existing = set(intext)
    for s in sup_citations:
        if s not in existing:
            intext.append(s)
            existing.add(s)
    return refs, intext


def extract_from_bibtex(path: Path) -> Tuple[List[ParsedReference], List[str]]:
    """Parse a BibTeX file into ParsedReference records.

    Per-entry parsing is wrapped so one malformed entry can't abort the file.
    """
    try:
        import bibtexparser
    except ImportError as e:
        raise RuntimeError(
            "bibtexparser is required to read .bib files. Install with: pip install 'bibtexparser<2'"
        ) from e
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            db = bibtexparser.load(fh)
    except Exception as e:
        raise RuntimeError(f"Could not parse BibTeX file {path.name}: {e}") from e
    refs: List[ParsedReference] = []
    for entry in db.entries:
        try:
            ref = ParsedReference(raw_text=str(entry))
            ref.title = (entry.get("title") or "").strip("{}").strip()
            ref.year = extract_year(str(entry.get("year", "")))
            ref.container_title = entry.get("journal") or entry.get("booktitle") or ""
            ref.volume = str(entry.get("volume", ""))
            ref.issue = str(entry.get("number", "") or entry.get("issue", ""))
            ref.pages = str(entry.get("pages", "")).replace("--", "-")
            ref.publisher = entry.get("publisher", "")
            ref.doi = extract_doi(entry.get("doi", "")) or (entry.get("doi") or "").lower().strip()
            ref.url = entry.get("url", "")
            ref.arxiv_id = extract_arxiv_id(entry.get("eprint", "") or entry.get("note", ""))
            if entry.get("author"):
                ref.authors = [Author.from_string(a) for a in entry["author"].split(" and ") if a.strip()]
            etype = (entry.get("ENTRYTYPE") or "").lower()
            ref.type = {
                "article": CitationType.JOURNAL_ARTICLE,
                "book": CitationType.BOOK,
                "incollection": CitationType.BOOK_CHAPTER,
                "inproceedings": CitationType.CONFERENCE_PAPER,
                "conference": CitationType.CONFERENCE_PAPER,
                "techreport": CitationType.REPORT,
                "phdthesis": CitationType.THESIS,
                "mastersthesis": CitationType.THESIS,
                "misc": CitationType.UNKNOWN,
            }.get(etype, CitationType.UNKNOWN)
            refs.append(ref)
        except Exception as e:  # pragma: no cover - defensive
            import logging
            logging.getLogger("biblio-check.extract").debug("bibtex entry failed: %s", e)
            refs.append(ParsedReference(raw_text=str(entry)))
    return refs, []


def extract_from_ris(path: Path) -> Tuple[List[ParsedReference], List[str]]:
    """Parse a RIS file. Per-record parsing wrapped for robustness."""
    try:
        import rispy
    except ImportError as e:
        raise RuntimeError(
            "rispy is required to read .ris files. Install with: pip install rispy"
        ) from e
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as fh:
            records = rispy.load(fh)
    except Exception as e:
        raise RuntimeError(f"Could not parse RIS file {path.name}: {e}") from e
    refs: List[ParsedReference] = []
    for rec in records:
        try:
            ref = ParsedReference(raw_text=str(rec))
            ref.title = rec.get("title") or rec.get("primary_title", "")
            ref.year = extract_year(str(rec.get("year", "")))
            ref.container_title = rec.get("journal_name") or rec.get("secondary_title", "")
            ref.volume = str(rec.get("volume", ""))
            ref.issue = str(rec.get("number", ""))
            sp, ep = rec.get("start_page", ""), rec.get("end_page", "")
            ref.pages = f"{sp}-{ep}" if sp and ep else (sp or "")
            ref.doi = (rec.get("doi") or "").lower()
            ref.publisher = rec.get("publisher", "")
            ref.url = rec.get("url", "")
            ref.authors = [Author.from_string(a) for a in rec.get("authors", []) if a]
            type_map = {
                "JOUR": CitationType.JOURNAL_ARTICLE,
                "BOOK": CitationType.BOOK,
                "CHAP": CitationType.BOOK_CHAPTER,
                "CONF": CitationType.CONFERENCE_PAPER,
                "CPAPER": CitationType.CONFERENCE_PAPER,
                "RPRT": CitationType.REPORT,
                "THES": CitationType.THESIS,
            }
            ref.type = type_map.get(rec.get("type_of_reference", ""), CitationType.UNKNOWN)
            refs.append(ref)
        except Exception as e:  # pragma: no cover - defensive
            import logging
            logging.getLogger("biblio-check.extract").debug("ris record failed: %s", e)
            refs.append(ParsedReference(raw_text=str(rec)))
    return refs, []


def extract(path: Path, parser_backend: str = "auto") -> Tuple[List[ParsedReference], List[str]]:
    """Dispatch on file extension. Raises RuntimeError with a clear message on
    unreadable files; never raises a bare traceback to the caller.

    `parser_backend` controls per-entry field parsing for text/docx/pdf inputs
    ('auto' uses AnyStyle/GROBID if installed, else regex). BibTeX/RIS are
    already structured, so the backend doesn't apply to them.
    """
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return extract_from_pdf(path, parser_backend=parser_backend)
        if suffix == ".docx":
            return extract_from_docx(path, parser_backend=parser_backend)
        if suffix == ".bib":
            return extract_from_bibtex(path)
        if suffix == ".ris":
            return extract_from_ris(path)
        # .txt, .md, unknown -> treat as text.
        return extract_from_text(
            path.read_text(encoding="utf-8", errors="replace"), parser_backend=parser_backend
        )
    except RuntimeError:
        raise  # already a clean, user-facing message
    except Exception as e:
        raise RuntimeError(
            f"Could not read {path.name} ({suffix or 'no extension'}): {e}. "
            f"Try converting it to a .txt file with one reference per line."
        ) from e
