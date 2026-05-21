"""Emit audit artifacts in the user's chosen format(s).

Output formats:
  * markdown  -- human-readable audit report
  * docx      -- annotated .docx with tracked-changes corrections
  * json      -- machine-readable, includes raw API evidence for audit
  * bib       -- corrected BibTeX
  * ris       -- corrected RIS

Tier display rules:
  * VERIFIED                     -> the citation is shown as-is, with a one-line confirmation.
  * VERIFIED_WITH_STYLE_EDITS    -> the citation is shown as-is, followed by an itemized list of
                                   cosmetic edits and the resulting corrected form.
  * PARTIAL                      -> a field-by-field diff table is shown, with severity, followed by
                                   the corrected form.
  * HALLUCINATED                 -> the original is shown with a clear warning that no canonical
                                   record was found.
"""
from __future__ import annotations

import json
import re
import datetime
from pathlib import Path
from typing import List, Optional, Dict

from .models import VerificationResult, Confidence
from .format_citation import format_reference, to_bibtex, SUPPORTED_STYLES


TIER_LABEL = {
    Confidence.VERIFIED: "VERIFIED",
    Confidence.VERIFIED_WITH_STYLE_EDITS: "VERIFIED (style edits suggested)",
    Confidence.VERIFIED_VIA_URL: "VERIFIED (grey literature / non-academic source)",
    Confidence.PARTIAL: "PARTIAL (substantive errors)",
    Confidence.HALLUCINATED: "HALLUCINATED (no canonical match)",
}


def _clean_bibliography_entries(results: List[VerificationResult], style: str) -> List[str]:
    """Build the ready-to-paste bibliography lines.

    For VERIFIED, VERIFIED_WITH_STYLE_EDITS, PARTIAL: emit the canonical
    formatted citation (consistent style across the whole list).
    For VERIFIED_VIA_URL (grey literature): keep the user's original wording
    -- we have no canonical metadata to reformat from, and the user's version
    is the authoritative one.
    For HALLUCINATED: emit a clearly-marked placeholder so the user sees
    something was removed and what it claimed to be.
    """
    out: List[str] = []
    for i, r in enumerate(results, start=1):
        if r.confidence == Confidence.HALLUCINATED:
            title_snippet = r.parsed.title or r.parsed.raw_text[:80] or "(no title)"
            out.append(f"[REMOVED — hallucinated entry could not be verified: \"{title_snippet}\"]")
        elif r.confidence == Confidence.VERIFIED_VIA_URL:
            # Grey literature: trust the user's wording verbatim.
            out.append(r.parsed.raw_text.strip())
        elif r.canonical:
            out.append(r.formatted.get(style) or format_reference(r.canonical, style, index=i))
        else:
            out.append(r.parsed.raw_text.strip())
    return out


def render_clean_bibliography(results: List[VerificationResult], style: str) -> str:
    """Plain-text clean bibliography for separate file output."""
    lines = _clean_bibliography_entries(results, style)
    return "\n\n".join(lines) + "\n"


def _render_duplicates_section(dups) -> List[str]:
    """Build the duplicate-entries section."""
    out: List[str] = []
    if not dups:
        return out
    out.append("## Duplicate bibliography entries")
    out.append("")
    out.append(f"Found {len(dups)} pair(s) of entries that appear to reference the same paper:")
    out.append("")
    for i, j, reason in dups:
        out.append(f"- Entries #{i} and #{j} -- {reason}.")
    out.append("")
    out.append("Consider keeping one and removing the other to clean up the bibliography.")
    out.append("")
    return out


# ---------- Markdown -------------------------------------------------------

def _entry_markdown(i: int, r: VerificationResult, style: str) -> List[str]:
    """Render one reference's section. Format depends on confidence tier."""
    out: List[str] = []
    label = TIER_LABEL[r.confidence]
    out.append(f"## {i}. {label}: {r.parsed.title or '(no title parsed)'}")
    out.append("")

    out.append("**As cited:**")
    out.append("")
    out.append(f"> {r.parsed.raw_text.strip() or '(no raw text)'}")
    out.append("")

    if r.confidence == Confidence.VERIFIED_VIA_URL:
        out.append("**Status:** This source is not indexed in academic databases "
                   "(CrossRef, OpenAlex, PubMed, Semantic Scholar, arXiv), which is normal for "
                   "grey literature: government reports, NGO publications, agency websites, "
                   "white papers, and working drafts. The cited URL resolves to a real page, "
                   "so the source legitimately exists. No corrections are proposed because "
                   "we cannot fetch authoritative metadata to compare against.")
        out.append("")
        if r.parsed.url:
            out.append(f"**Cited URL:** [{r.parsed.url}]({r.parsed.url}) -- HEAD-verified to resolve.")
            out.append("")
        out.append("Keep as written. If the source is critical to your argument, manually open "
                   "the URL and confirm the title, date, and authoring organization.")
        out.append("")

    elif r.confidence == Confidence.VERIFIED and r.canonical:
        out.append("**Status:** All substantive components match the canonical record. "
                   "No changes are needed.")
        out.append("")
        out.append(f"**Canonical record:** {r.canonical.source.upper()} -> {r.canonical.source_url}")
        if r.canonical.doi:
            out.append(f"- DOI: [{r.canonical.doi}](https://doi.org/{r.canonical.doi})")
        if r.canonical.pmid:
            out.append(f"- PMID: [{r.canonical.pmid}](https://pubmed.ncbi.nlm.nih.gov/{r.canonical.pmid}/)")
        if r.canonical.arxiv_id:
            out.append(f"- arXiv: [{r.canonical.arxiv_id}](https://arxiv.org/abs/{r.canonical.arxiv_id})")
        out.append("")

    elif r.confidence == Confidence.VERIFIED_WITH_STYLE_EDITS and r.canonical:
        out.append("**Status:** All substantive components are correct. The following cosmetic edits "
                   f"would bring the entry into strict {style.upper()} style:")
        out.append("")
        for j, e in enumerate(r.style_edits, start=1):
            out.append(f"{j}. {e.description}")
        out.append("")
        out.append(f"**Canonical record:** {r.canonical.source.upper()} -> {r.canonical.source_url}")
        if r.canonical.doi:
            out.append(f"- DOI: [{r.canonical.doi}](https://doi.org/{r.canonical.doi})")
        out.append("")
        out.append(f"**With all style edits applied ({style.upper()}):**")
        out.append("")
        corrected = r.formatted.get(style) or format_reference(r.canonical, style, index=i)
        out.append(f"> {corrected}")
        out.append("")

    elif r.confidence == Confidence.PARTIAL and r.canonical:
        out.append("**Status:** A real paper matching most fields was found, but the citation as "
                   "written has substantive errors. Field-by-field audit:")
        out.append("")
        out.append("| Field | As cited | Canonical | Severity |")
        out.append("|---|---|---|---|")
        for d in r.discrepancies:
            cited = (d.parsed_value or "(missing)").replace("|", "\\|")
            canon = (d.canonical_value or "(missing)").replace("|", "\\|")
            out.append(f"| {d.field} | {cited} | {canon} | {d.severity} |")
        out.append("")
        if r.style_edits:
            out.append("Additional cosmetic style edits also recommended:")
            for j, e in enumerate(r.style_edits, start=1):
                out.append(f"{j}. {e.description}")
            out.append("")
        out.append(f"**Canonical record:** {r.canonical.source.upper()} -> {r.canonical.source_url}")
        if r.canonical.doi:
            out.append(f"- DOI: [{r.canonical.doi}](https://doi.org/{r.canonical.doi})")
        if r.canonical.pmid:
            out.append(f"- PMID: [{r.canonical.pmid}](https://pubmed.ncbi.nlm.nih.gov/{r.canonical.pmid}/)")
        if r.canonical.arxiv_id:
            out.append(f"- arXiv: [{r.canonical.arxiv_id}](https://arxiv.org/abs/{r.canonical.arxiv_id})")
        out.append("")
        out.append(f"**Proposed corrected citation ({style.upper()}):**")
        out.append("")
        corrected = r.formatted.get(style) or format_reference(r.canonical, style, index=i)
        out.append(f"> {corrected}")
        out.append("")

    else:  # HALLUCINATED
        out.append("**Status:** No verification API returned a record that confidently matches this citation.")
        out.append("")
        out.append("APIs queried: CrossRef, OpenAlex, PubMed, Semantic Scholar, arXiv. ")
        out.append("This is the pattern typical of fabricated citations. ")
        out.append("Confirm independently or remove from the bibliography.")
        out.append("")
        if r.canonical:
            out.append(
                f"The closest candidate found was:")
            out.append("")
            out.append(f"> {r.canonical.title} ({r.canonical.year or 'n.d.'}) "
                       f"-- {r.canonical.source_url}")
            out.append("")
            out.append("This was below the confidence threshold and is shown for reference only; do NOT cite it as the user's intended source.")
            out.append("")

    if r.notes:
        out.append("**Notes:**")
        for nt in r.notes:
            out.append(f"- {nt}")
        out.append("")

    # Corroborating sources are useful evidence regardless of tier.
    corroborating = [m for m in r.matches if m is not r.canonical] if r.canonical else []
    if corroborating:
        out.append("**Cross-verified by:** " +
                   ", ".join(sorted({m.source for m in corroborating})))
        out.append("")

    out.append("---")
    out.append("")
    return out


def _render_crosscheck_section(crosscheck) -> List[str]:
    """Build the in-text vs bibliography cross-check section."""
    out: List[str] = []
    out.append("## In-text vs bibliography cross-check")
    out.append("")
    n_orphans = len(crosscheck.orphan_intext)
    n_unused = len(crosscheck.unused_bibliography)
    if not n_orphans and not n_unused:
        out.append("Every in-text citation has a matching bibliography entry, and every "
                   "bibliography entry is cited at least once in the body. No orphans or unused entries.")
        out.append("")
        return out

    if n_orphans:
        out.append(f"### Orphan in-text citations ({n_orphans})")
        out.append("")
        out.append("These in-text citations appear in the body text but no matching entry was "
                   "found in the bibliography. Common LLM failure mode: the body cites a paper "
                   "that was never added to the references.")
        out.append("")
        for c in crosscheck.orphan_intext:
            out.append(f"- `{c.raw}`")
        out.append("")
    if n_unused:
        out.append(f"### Unused bibliography entries ({n_unused})")
        out.append("")
        out.append("These bibliography entries are never cited in the body text. Either remove "
                   "them or add the missing in-text citation.")
        out.append("")
        for idx in crosscheck.unused_bibliography:
            out.append(f"- Entry #{idx}")
        out.append("")
    out.append("Note: this cross-check is based on simple author+year or numeric pattern matching. "
               "Citations using unusual formats may be flagged incorrectly. Use judgment.")
    out.append("")
    return out


def render_markdown(results: List[VerificationResult], style: str, source_name: str = "",
                    crosscheck=None, duplicates=None) -> str:
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    n = len(results)
    counts: Dict[Confidence, int] = {k: 0 for k in TIER_LABEL}
    for r in results:
        counts[r.confidence] = counts.get(r.confidence, 0) + 1

    out: List[str] = []
    out.append(f"# Bibliography audit ({style.upper()})")
    out.append("")
    if source_name:
        out.append(f"Source: `{source_name}`  ")
    out.append(f"Generated: {now}  ")
    out.append(f"References checked: {n}")
    out.append("")
    out.append("## Summary")
    out.append("")
    out.append("| Tier | Count | Meaning |")
    out.append("|---|---|---|")
    out.append(f"| VERIFIED | {counts[Confidence.VERIFIED]} | All components correct, no changes needed. |")
    out.append(f"| VERIFIED (style edits) | {counts[Confidence.VERIFIED_WITH_STYLE_EDITS]} | Content correct, cosmetic polish suggested. |")
    out.append(f"| VERIFIED (grey literature) | {counts[Confidence.VERIFIED_VIA_URL]} | Not in academic databases; URL resolves. Government/NGO/agency sources. |")
    out.append(f"| PARTIAL | {counts[Confidence.PARTIAL]} | Real paper but citation has substantive errors. |")
    out.append(f"| HALLUCINATED | {counts[Confidence.HALLUCINATED]} | No canonical record found. |")
    out.append("")

    # Order: hallucinated first, then partial, then verified-with-edits, then verified (grey), then verified.
    order = {
        Confidence.HALLUCINATED: 0,
        Confidence.PARTIAL: 1,
        Confidence.VERIFIED_WITH_STYLE_EDITS: 2,
        Confidence.VERIFIED_VIA_URL: 3,
        Confidence.VERIFIED: 4,
    }
    indexed = list(enumerate(results, start=1))
    indexed.sort(key=lambda p: order.get(p[1].confidence, 99))

    for i, r in indexed:
        out.extend(_entry_markdown(i, r, style))

    out.append("## How to read this report")
    out.append("")
    out.append("- **VERIFIED**: two or more independent metadata APIs returned matching records "
               "for every substantive field. The citation can stay as written.")
    out.append("- **VERIFIED (style edits)**: substantive content is right; the suggestions are "
               "cosmetic conformity with the chosen style and do not change meaning.")
    out.append("- **VERIFIED (grey literature)**: not in any academic database, but either the "
               "cited URL HEAD-resolved to a real page or the entry is structurally clearly "
               "institutional (government / NGO / agency report with a named publisher). "
               "Adding a URL to citations of this kind lets future runs verify them more directly.")
    out.append("- **PARTIAL**: a real paper exists but the cited form contains content errors "
               "(wrong year, wrong author, wrong DOI, etc.). The corrected form is provided.")
    out.append("- **HALLUCINATED**: no API returned a record matching the citation. Treat as "
               "fabricated until manually confirmed.")
    out.append("")
    out.append("APIs queried: CrossRef, OpenAlex, PubMed, Semantic Scholar, arXiv. ")
    out.append("Verification confirms existence and metadata accuracy. It does NOT confirm that "
               "the cited work actually supports the claim it is attached to in the body text -- "
               "that semantic check is the author's responsibility.")
    out.append("")

    # In-text vs bibliography cross-check (if a paper body was parsed).
    if crosscheck is not None:
        out.extend(_render_crosscheck_section(crosscheck))

    # Duplicate bibliography entries.
    if duplicates:
        out.extend(_render_duplicates_section(duplicates))

    # Clean bibliography: ready-to-paste, in original order.
    out.append(f"## Clean bibliography ({style.upper()}) -- ready to paste")
    out.append("")
    out.append(
        "Verified academic entries are re-emitted in canonical form for style consistency. "
        "Grey-literature entries are kept verbatim (we have no authoritative metadata to reformat from). "
        "Hallucinated entries appear as bracketed placeholders so you can see what was removed."
    )
    out.append("")
    for entry in _clean_bibliography_entries(results, style):
        out.append(entry)
        out.append("")
    return "\n".join(out)


# ---------- HTML ------------------------------------------------------------

_HTML_TIER_COLOR = {
    Confidence.VERIFIED: "#0a7a0a",
    Confidence.VERIFIED_WITH_STYLE_EDITS: "#946600",
    Confidence.VERIFIED_VIA_URL: "#1f57c2",
    Confidence.PARTIAL: "#bd5d00",
    Confidence.HALLUCINATED: "#b00020",
}


def _html_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
         .replace('"', "&quot;")
    )


def _html_entry(i: int, r: VerificationResult, style: str) -> str:
    color = _HTML_TIER_COLOR.get(r.confidence, "#444")
    tier_label = _html_escape(TIER_LABEL[r.confidence])
    title = _html_escape(r.parsed.title or "(no title parsed)")
    raw = _html_escape(r.parsed.raw_text.strip() or "(no raw text)")

    sections = [
        f'<details class="entry" data-tier="{r.confidence.value}" open>',
        f'  <summary class="tier-{r.confidence.value}">'
        f'<span class="badge" style="background:{color}">{tier_label}</span>'
        f' <span class="entry-num">#{i}</span> {title}</summary>',
        '  <div class="entry-body">',
        f'    <div class="ascited"><b>As cited:</b><blockquote>{raw}</blockquote></div>',
    ]

    if r.confidence == Confidence.VERIFIED_VIA_URL:
        if r.parsed.url:
            url_esc = _html_escape(r.parsed.url)
            sections.append(
                f'    <p><b>Status:</b> grey literature. URL HEAD-verified: '
                f'<a href="{url_esc}" target="_blank">{url_esc}</a></p>'
            )
        else:
            sections.append(
                '    <p><b>Status:</b> grey literature; institutional author + named publisher detected.</p>'
            )

    elif r.confidence == Confidence.VERIFIED and r.canonical:
        sections.append('    <p><b>Status:</b> all components verified. No changes needed.</p>')
        if r.canonical.doi:
            sections.append(
                f'    <p><b>DOI:</b> <a href="https://doi.org/{_html_escape(r.canonical.doi)}" '
                f'target="_blank">{_html_escape(r.canonical.doi)}</a></p>'
            )

    elif r.confidence == Confidence.VERIFIED_WITH_STYLE_EDITS and r.canonical:
        sections.append('    <p><b>Status:</b> content correct; cosmetic edits suggested:</p>')
        sections.append('    <ul>')
        for e in r.style_edits:
            sections.append(f'      <li>{_html_escape(e.description)}</li>')
        sections.append('    </ul>')
        corrected = r.formatted.get(style) or format_reference(r.canonical, style, index=i)
        sections.append(
            f'    <p><b>Corrected ({style.upper()}):</b></p>'
            f'    <blockquote class="corrected">{_html_escape(corrected)}</blockquote>'
        )

    elif r.confidence == Confidence.PARTIAL and r.canonical:
        sections.append('    <p><b>Status:</b> real paper found; citation has substantive errors:</p>')
        sections.append('    <table class="diff"><thead><tr><th>Field</th><th>As cited</th>'
                        '<th>Canonical</th><th>Severity</th></tr></thead><tbody>')
        for d in r.discrepancies:
            cls = "sev-major" if d.severity == "major" else "sev-minor"
            sections.append(
                f'      <tr class="{cls}">'
                f'<td>{_html_escape(d.field)}</td>'
                f'<td>{_html_escape(d.parsed_value or "(missing)")}</td>'
                f'<td>{_html_escape(d.canonical_value or "(missing)")}</td>'
                f'<td>{_html_escape(d.severity)}</td></tr>'
            )
        sections.append('    </tbody></table>')
        if r.style_edits:
            sections.append('    <p>Additional style edits:</p><ul>')
            for e in r.style_edits:
                sections.append(f'      <li>{_html_escape(e.description)}</li>')
            sections.append('    </ul>')
        corrected = r.formatted.get(style) or format_reference(r.canonical, style, index=i)
        sections.append(
            f'    <p><b>Proposed corrected ({style.upper()}):</b></p>'
            f'    <blockquote class="corrected">{_html_escape(corrected)}</blockquote>'
        )

    else:  # HALLUCINATED
        sections.append(
            '    <p><b>Status:</b> no verification API returned a confident match. '
            'Treat as fabricated until manually confirmed.</p>'
        )
        if r.canonical:
            t = _html_escape(r.canonical.title or "")
            u = _html_escape(r.canonical.source_url or "")
            sections.append(
                f'    <p>Closest candidate (for context only, do NOT cite as the user\'s intended source):'
                f' <a href="{u}" target="_blank">{t}</a></p>'
            )

    if r.notes:
        sections.append('    <details class="notes"><summary>Notes</summary><ul>')
        for nt in r.notes:
            sections.append(f'      <li>{_html_escape(nt)}</li>')
        sections.append('    </ul></details>')

    corroborating = [m for m in r.matches if m is not r.canonical] if r.canonical else []
    if corroborating:
        srcs = sorted({m.source for m in corroborating})
        sections.append(
            f'    <p class="corroborated"><b>Cross-verified by:</b> {", ".join(srcs)}</p>'
        )

    sections.append('  </div>')
    sections.append('</details>')
    return "\n".join(sections)


def render_html(results: List[VerificationResult], style: str, source_name: str = "",
                crosscheck=None, duplicates=None) -> str:
    """Single-file HTML report with color-coded tiers, collapsible entries,
    clickable links, and a copy-to-clipboard clean-bibliography section."""
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    n = len(results)
    counts: Dict[Confidence, int] = {k: 0 for k in TIER_LABEL}
    for r in results:
        counts[r.confidence] = counts.get(r.confidence, 0) + 1

    css = """
      body { font: 15px/1.45 -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
             max-width: 980px; margin: 2em auto; padding: 0 1em; color: #222; }
      h1 { margin-bottom: 0.2em; }
      .meta { color: #666; font-size: 0.92em; margin-bottom: 1.5em; }
      table.summary, table.diff { border-collapse: collapse; width: 100%; margin: 0.6em 0 1em; }
      table.summary th, table.summary td, table.diff th, table.diff td {
        border: 1px solid #ddd; padding: 6px 10px; text-align: left; vertical-align: top;
      }
      table.summary th { background: #f6f8fa; }
      .badge { display: inline-block; color: white; padding: 1px 8px; border-radius: 10px;
               font-size: 0.78em; font-weight: 600; margin-right: 6px; }
      details.entry { margin: 0.4em 0; border-left: 4px solid #aaa; padding-left: 12px; }
      details.entry[data-tier="verified"] { border-left-color: #0a7a0a; }
      details.entry[data-tier="verified_with_style_edits"] { border-left-color: #946600; }
      details.entry[data-tier="verified_via_url"] { border-left-color: #1f57c2; }
      details.entry[data-tier="partial"] { border-left-color: #bd5d00; }
      details.entry[data-tier="hallucinated"] { border-left-color: #b00020; }
      details.entry summary { cursor: pointer; padding: 4px 0; }
      details.entry summary::-webkit-details-marker { color: #888; }
      .entry-num { color: #888; font-weight: 600; }
      .entry-body { margin: 0.5em 0 0.8em 0; padding-left: 8px; }
      blockquote { background: #f6f8fa; border-left: 3px solid #ccc; padding: 8px 12px;
                   margin: 0.4em 0; font-family: ui-monospace, SFMono-Regular, Menlo, monospace;
                   font-size: 0.92em; white-space: pre-wrap; word-break: break-word; }
      blockquote.corrected { border-left-color: #1f57c2; background: #eef4fc; }
      tr.sev-major { background: #fdecec; }
      tr.sev-minor { background: #fef6e7; }
      details.notes { margin: 0.4em 0 0.6em; }
      details.notes summary { color: #555; font-size: 0.92em; }
      .corroborated { color: #555; font-size: 0.9em; }
      .clean-bib-entry { padding: 4px 0; }
      .clean-bib-entry.removed { color: #b00020; font-style: italic; }
      button.copy { font: inherit; padding: 4px 10px; cursor: pointer; }
      .filters { margin: 1em 0; }
      .filters label { margin-right: 12px; font-size: 0.92em; cursor: pointer; }
      .crosscheck-section, .duplicates-section { margin-top: 1.5em; padding: 0.5em 1em;
                                                  background: #f6f8fa; border-radius: 4px; }
    """

    js = r"""
      function copyCleanBib() {
        const el = document.getElementById('clean-bib');
        const text = Array.from(el.querySelectorAll('.clean-bib-entry')).map(e => e.innerText).join('\n\n');
        navigator.clipboard.writeText(text).then(() => {
          const btn = document.getElementById('copy-btn');
          const old = btn.innerText;
          btn.innerText = 'Copied!';
          setTimeout(() => { btn.innerText = old; }, 1200);
        });
      }
      function applyFilter() {
        const show = {};
        document.querySelectorAll('.filters input').forEach(cb => { show[cb.dataset.tier] = cb.checked; });
        document.querySelectorAll('details.entry').forEach(d => {
          d.style.display = show[d.dataset.tier] ? '' : 'none';
        });
      }
    """

    order = {
        Confidence.HALLUCINATED: 0,
        Confidence.PARTIAL: 1,
        Confidence.VERIFIED_WITH_STYLE_EDITS: 2,
        Confidence.VERIFIED_VIA_URL: 3,
        Confidence.VERIFIED: 4,
    }
    indexed = sorted(enumerate(results, start=1), key=lambda p: order.get(p[1].confidence, 99))

    parts: List[str] = []
    parts.append("<!doctype html><html><head><meta charset='utf-8'>")
    parts.append(f"<title>Bibliography audit ({_html_escape(style.upper())})</title>")
    parts.append(f"<style>{css}</style></head><body>")
    parts.append(f"<h1>Bibliography audit ({_html_escape(style.upper())})</h1>")
    if source_name:
        parts.append(f"<div class='meta'>Source: <code>{_html_escape(source_name)}</code></div>")
    parts.append(f"<div class='meta'>Generated: {now} &middot; References checked: {n}</div>")

    # Summary table
    parts.append("<h2>Summary</h2>")
    parts.append("<table class='summary'><thead><tr><th>Tier</th><th>Count</th><th>Meaning</th></tr></thead><tbody>")
    rows = [
        ("verified", "VERIFIED", "All components correct, no changes needed.", counts[Confidence.VERIFIED]),
        ("verified_with_style_edits", "VERIFIED (style edits)", "Content correct, cosmetic polish suggested.", counts[Confidence.VERIFIED_WITH_STYLE_EDITS]),
        ("verified_via_url", "VERIFIED (grey literature)", "Not in academic databases; URL resolves OR institutional source.", counts[Confidence.VERIFIED_VIA_URL]),
        ("partial", "PARTIAL", "Real paper but citation has substantive errors.", counts[Confidence.PARTIAL]),
        ("hallucinated", "HALLUCINATED", "No canonical record found.", counts[Confidence.HALLUCINATED]),
    ]
    for tier_key, label, meaning, ct in rows:
        color = _HTML_TIER_COLOR.get(getattr(Confidence, tier_key.upper()), "#444")
        parts.append(
            f"<tr><td><span class='badge' style='background:{color}'>{label}</span></td>"
            f"<td>{ct}</td><td>{meaning}</td></tr>"
        )
    parts.append("</tbody></table>")

    # Filter controls
    parts.append("<div class='filters'><b>Show:</b> ")
    for tier_key, label, _meaning, _ct in rows:
        parts.append(
            f"<label><input type='checkbox' data-tier='{tier_key}' checked onchange='applyFilter()'> {label}</label>"
        )
    parts.append("</div>")

    parts.append("<h2>References</h2>")
    for i, r in indexed:
        parts.append(_html_entry(i, r, style))

    # Cross-check section
    if crosscheck is not None:
        parts.append("<section class='crosscheck-section'>")
        parts.append("<h2>In-text vs bibliography cross-check</h2>")
        n_orphans = len(crosscheck.orphan_intext)
        n_unused = len(crosscheck.unused_bibliography)
        if not n_orphans and not n_unused:
            parts.append("<p>Every in-text citation has a matching bibliography entry, and every "
                         "bibliography entry is cited at least once.</p>")
        else:
            if n_orphans:
                parts.append(f"<p><b>Orphan in-text citations ({n_orphans}):</b></p><ul>")
                for c in crosscheck.orphan_intext:
                    parts.append(f"<li><code>{_html_escape(c.raw)}</code></li>")
                parts.append("</ul>")
            if n_unused:
                parts.append(f"<p><b>Unused bibliography entries ({n_unused}):</b></p><ul>")
                for idx in crosscheck.unused_bibliography:
                    parts.append(f"<li>Entry #{idx}</li>")
                parts.append("</ul>")
        parts.append("</section>")

    # Duplicates section
    if duplicates:
        parts.append("<section class='duplicates-section'>")
        parts.append("<h2>Duplicate bibliography entries</h2><ul>")
        for i, j, reason in duplicates:
            parts.append(f"<li>Entries #{i} and #{j} -- {_html_escape(reason)}.</li>")
        parts.append("</ul></section>")

    # Clean bibliography with copy-to-clipboard
    parts.append("<h2>Clean bibliography ({}) <button id='copy-btn' class='copy' onclick='copyCleanBib()'>Copy all</button></h2>".format(_html_escape(style.upper())))
    parts.append("<div id='clean-bib'>")
    for entry in _clean_bibliography_entries(results, style):
        cls = "clean-bib-entry removed" if entry.startswith("[REMOVED") else "clean-bib-entry"
        parts.append(f"<div class='{cls}'>{_html_escape(entry)}</div>")
    parts.append("</div>")

    parts.append(f"<script>{js}</script>")
    parts.append("</body></html>")
    return "\n".join(parts)


# ---------- JSON ------------------------------------------------------------

def render_json(results: List[VerificationResult]) -> str:
    return json.dumps([r.to_dict() for r in results], indent=2, ensure_ascii=False)


# ---------- BibTeX ----------------------------------------------------------

def render_bibtex(results: List[VerificationResult]) -> str:
    entries = []
    for i, r in enumerate(results, start=1):
        if r.confidence == Confidence.VERIFIED_VIA_URL:
            # Grey literature: emit a @misc entry from the parsed fields.
            p = r.parsed
            first_surname = (p.authors[0].family if p.authors else "anon").lower()
            key = re.sub(r"[^a-z0-9]+", "", first_surname) + str(p.year or "nd")
            lines = [f"@misc{{{key},"]
            if p.authors:
                lines.append(f"  author = {{{ ' and '.join(a.full for a in p.authors) }}},")
            if p.title:
                lines.append(f"  title = {{{p.title}}},")
            if p.year:
                lines.append(f"  year = {{{p.year}}},")
            if p.publisher:
                lines.append(f"  publisher = {{{p.publisher}}},")
            if p.url:
                lines.append(f"  url = {{{p.url}}},")
            lines.append("  note = {Grey literature; verified by URL resolution only}")
            lines.append("}")
            entries.append("\n".join(lines))
        elif r.canonical and r.confidence != Confidence.HALLUCINATED:
            entries.append(to_bibtex(r.canonical))
        else:
            entries.append(
                f"% HALLUCINATED -- not emitting canonical BibTeX for: "
                f"{r.parsed.raw_text[:120]}"
            )
    return "\n\n".join(entries)


# ---------- RIS -------------------------------------------------------------

def _ris_for(c) -> str:
    """Emit one RIS record."""
    if not c:
        return ""
    type_map = {
        "journal-article": "JOUR",
        "book": "BOOK",
        "book-chapter": "CHAP",
        "proceedings-article": "CONF",
        "posted-content": "JOUR",
        "report": "RPRT",
        "thesis": "THES",
    }
    rtype = type_map.get(c.type.value if hasattr(c.type, "value") else str(c.type), "GEN")
    lines = [f"TY  - {rtype}"]
    for a in c.authors:
        lines.append(f"AU  - {a.family}, {a.given}".rstrip(", "))
    if c.title:
        lines.append(f"TI  - {c.title}")
    if c.year:
        lines.append(f"PY  - {c.year}")
    if c.container_title:
        lines.append(f"JO  - {c.container_title}")
    if c.volume:
        lines.append(f"VL  - {c.volume}")
    if c.issue:
        lines.append(f"IS  - {c.issue}")
    if c.pages:
        if "-" in c.pages or "–" in c.pages:
            sp, _, ep = c.pages.replace("–", "-").partition("-")
            lines.append(f"SP  - {sp.strip()}")
            lines.append(f"EP  - {ep.strip()}")
        else:
            lines.append(f"SP  - {c.pages}")
    if c.doi:
        lines.append(f"DO  - {c.doi}")
    if c.publisher:
        lines.append(f"PB  - {c.publisher}")
    lines.append("ER  - ")
    return "\n".join(lines)


def render_ris(results: List[VerificationResult]) -> str:
    out: List[str] = []
    for r in results:
        if r.confidence == Confidence.VERIFIED_VIA_URL:
            # Grey literature: emit a generic record from parsed fields.
            p = r.parsed
            lines = ["TY  - GEN"]
            for a in p.authors:
                lines.append(f"AU  - {a.family}, {a.given}".rstrip(", "))
            if p.title:
                lines.append(f"TI  - {p.title}")
            if p.year:
                lines.append(f"PY  - {p.year}")
            if p.publisher:
                lines.append(f"PB  - {p.publisher}")
            if p.url:
                lines.append(f"UR  - {p.url}")
            lines.append("N1  - Grey literature; verified by URL resolution only.")
            lines.append("ER  - ")
            out.append("\n".join(lines))
        elif r.canonical and r.confidence != Confidence.HALLUCINATED:
            out.append(_ris_for(r.canonical))
        else:
            out.append(
                f"% HALLUCINATED -- not emitting canonical RIS for: "
                f"{r.parsed.raw_text[:120]}"
            )
    return "\n\n".join(out)


# ---------- Annotated .docx with tracked changes ---------------------------

def render_annotated_docx(
    results: List[VerificationResult],
    style: str,
    out_path: Path,
    source_name: str = "",
    crosscheck=None,
    duplicates=None,
) -> None:
    """Write a .docx where each reference shows the original wording with
    corrections as Word tracked changes.

    Behaviour per tier:
      * VERIFIED: shown as-is, with a comment that no edits are needed.
      * VERIFIED with style edits: tracked-change conversion from original to
        corrected form, with an audit-trail line beneath listing the edits.
      * PARTIAL: tracked-change conversion AND a separate audit table showing
        every substantive diff with severity.
      * HALLUCINATED: original kept, with bold red warning, no insertion.
    """
    import docx
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()
    now_iso = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    author = "Editor"

    doc.add_heading(f"Bibliography audit ({style.upper()})", level=0)
    if source_name:
        doc.add_paragraph(f"Source: {source_name}")
    doc.add_paragraph(f"Generated: {now_iso}")
    doc.add_paragraph(f"References checked: {len(results)}")

    # Summary table
    summary = doc.add_paragraph()
    summary.add_run("Summary: ").bold = True
    counts = {k: 0 for k in TIER_LABEL}
    for r in results:
        counts[r.confidence] = counts.get(r.confidence, 0) + 1
    summary.add_run(
        f"VERIFIED = {counts[Confidence.VERIFIED]}, "
        f"VERIFIED with style edits = {counts[Confidence.VERIFIED_WITH_STYLE_EDITS]}, "
        f"VERIFIED (grey literature) = {counts[Confidence.VERIFIED_VIA_URL]}, "
        f"PARTIAL = {counts[Confidence.PARTIAL]}, "
        f"HALLUCINATED = {counts[Confidence.HALLUCINATED]}."
    )

    doc.add_heading("References", level=1)

    def _add_tracked_change(paragraph, original: str, corrected: str, rev_id: int) -> None:
        """Insert a w:del + w:ins pair so Word sees a real tracked edit."""
        del_elem = OxmlElement("w:del")
        del_elem.set(qn("w:id"), str(rev_id))
        del_elem.set(qn("w:author"), author)
        del_elem.set(qn("w:date"), now_iso)
        r = OxmlElement("w:r")
        r.append(OxmlElement("w:rPr"))
        del_text = OxmlElement("w:delText")
        del_text.set(qn("xml:space"), "preserve")
        del_text.text = original
        r.append(del_text)
        del_elem.append(r)
        paragraph._p.append(del_elem)

        ins_elem = OxmlElement("w:ins")
        ins_elem.set(qn("w:id"), str(rev_id + 1))
        ins_elem.set(qn("w:author"), author)
        ins_elem.set(qn("w:date"), now_iso)
        r2 = OxmlElement("w:r")
        r2.append(OxmlElement("w:rPr"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = corrected
        r2.append(t)
        ins_elem.append(r2)
        paragraph._p.append(ins_elem)

    rev_id = 1000
    for i, r in enumerate(results, start=1):
        # Header line: "N. TIER: title"
        header = doc.add_paragraph()
        run = header.add_run(f"{i}. {TIER_LABEL[r.confidence]}")
        run.bold = True
        if r.confidence == Confidence.HALLUCINATED:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif r.confidence == Confidence.PARTIAL:
            run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
        elif r.confidence == Confidence.VERIFIED_WITH_STYLE_EDITS:
            run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
        elif r.confidence == Confidence.VERIFIED_VIA_URL:
            run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)  # blue
        else:
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

        body = doc.add_paragraph(style="List Number")
        original = r.parsed.raw_text.strip()
        corrected = None
        if r.canonical and r.confidence != Confidence.HALLUCINATED:
            corrected = r.formatted.get(style) or format_reference(r.canonical, style, index=i)

        if r.confidence == Confidence.VERIFIED:
            body.add_run(original)
            note = doc.add_paragraph()
            note_run = note.add_run("  -> All components verified; no changes needed.")
            note_run.italic = True
            note_run.font.size = Pt(10)

        elif r.confidence == Confidence.VERIFIED_VIA_URL:
            body.add_run(original)
            note = doc.add_paragraph()
            note_run = note.add_run(
                "  -> Grey literature (government/NGO/agency source). Not indexed in academic "
                "databases, but the cited URL resolves to a real page. Keep as written."
            )
            note_run.italic = True
            note_run.font.size = Pt(10)

        elif r.confidence in (Confidence.VERIFIED_WITH_STYLE_EDITS, Confidence.PARTIAL):
            if corrected and original != corrected:
                _add_tracked_change(body, original, corrected, rev_id)
                rev_id += 2
            else:
                body.add_run(corrected or original)
            audit = doc.add_paragraph()
            audit_run = audit.add_run("  Audit trail:")
            audit_run.italic = True
            audit_run.font.size = Pt(10)
            for d in r.discrepancies:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{d.field}: cited '{d.parsed_value or '(missing)'}' but canonical "
                          f"is '{d.canonical_value or '(missing)'}' ({d.severity}).")
            for e in r.style_edits:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(e.description)

        else:  # HALLUCINATED
            run2 = body.add_run(original or "(no parsed text)")
            run2.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            warn = doc.add_paragraph()
            wrun = warn.add_run("  WARNING: No verification API returned a confident match. "
                                "Treat as fabricated until manually confirmed.")
            wrun.bold = True
            wrun.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # In-text vs bibliography cross-check section.
    if crosscheck is not None:
        doc.add_heading("In-text vs bibliography cross-check", level=1)
        n_orphans = len(crosscheck.orphan_intext)
        n_unused = len(crosscheck.unused_bibliography)
        if not n_orphans and not n_unused:
            doc.add_paragraph(
                "Every in-text citation has a matching bibliography entry, and every "
                "bibliography entry is cited at least once. No orphans or unused entries."
            )
        else:
            if n_orphans:
                p = doc.add_paragraph()
                p.add_run(f"Orphan in-text citations ({n_orphans}): ").bold = True
                p.add_run("cited in the body text but no matching bibliography entry was found.")
                for c in crosscheck.orphan_intext:
                    doc.add_paragraph(c.raw, style="List Bullet")
            if n_unused:
                p = doc.add_paragraph()
                p.add_run(f"Unused bibliography entries ({n_unused}): ").bold = True
                p.add_run("listed in the bibliography but never cited in the body.")
                for idx in crosscheck.unused_bibliography:
                    doc.add_paragraph(f"Entry #{idx}", style="List Bullet")

    # Duplicate bibliography entries.
    if duplicates:
        doc.add_heading("Duplicate bibliography entries", level=1)
        doc.add_paragraph(
            f"Found {len(duplicates)} pair(s) of entries that appear to reference the same paper."
        )
        for i, j, reason in duplicates:
            doc.add_paragraph(f"Entries #{i} and #{j} -- {reason}.", style="List Bullet")

    # Clean bibliography section (ready to paste).
    doc.add_heading(f"Clean bibliography ({style.upper()}) -- ready to paste", level=1)
    doc.add_paragraph(
        "Verified academic entries are re-emitted in canonical form for style consistency. "
        "Grey-literature entries are kept verbatim (we have no authoritative metadata to reformat from). "
        "Hallucinated entries appear as bracketed placeholders so you can see what was removed."
    )
    for entry in _clean_bibliography_entries(results, style):
        # Mark hallucinated placeholders with red text so they stand out.
        para = doc.add_paragraph()
        if entry.startswith("[REMOVED"):
            run = para.add_run(entry)
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.bold = True
        else:
            para.add_run(entry)

    doc.save(str(out_path))


# ---------- Top-level writer -----------------------------------------------

# Which formats to default to, by source extension.
DEFAULT_FORMAT_BY_INPUT = {
    ".docx": ["docx"],
    ".bib": ["bibtex"],
    ".ris": ["ris"],
    ".pdf": ["docx"],
    ".txt": ["docx"],
    ".md": ["docx"],
    "": ["docx"],
}


def write_all(
    results: List[VerificationResult],
    style: str,
    out_dir: Path,
    source_name: str = "",
    formats: Optional[List[str]] = None,
    crosscheck=None,
    duplicates=None,
) -> Dict[str, str]:
    """Write the requested output formats. Returns paths of files written.

    `formats` is a list of any of: 'markdown', 'json', 'docx', 'bibtex', 'ris'.
    If None, defaults to all five.
    """
    if formats is None:
        formats = ["markdown", "json", "docx", "bibtex", "ris"]

    # Populate formatted strings so markdown / docx / bibtex agree on output.
    for i, r in enumerate(results, start=1):
        if r.canonical:
            r.formatted[style] = format_reference(r.canonical, style, index=i)

    out_dir.mkdir(parents=True, exist_ok=True)
    written: Dict[str, str] = {}

    def _safe_write(fmt_key: str, filename: str, render_fn) -> None:
        """Render and write one format; record an error key instead of raising
        so one failing format never blocks the others."""
        try:
            content = render_fn()
            p = out_dir / filename
            p.write_text(content, encoding="utf-8")
            written[fmt_key] = str(p)
        except Exception as e:  # pragma: no cover - defensive
            import logging
            logging.getLogger("biblio-check.report").exception("%s render failed", fmt_key)
            written[f"{fmt_key}_error"] = str(e)

    if "markdown" in formats:
        _safe_write("markdown", "audit.md", lambda: render_markdown(
            results, style=style, source_name=source_name,
            crosscheck=crosscheck, duplicates=duplicates))

    if "json" in formats:
        _safe_write("json", "audit.json", lambda: render_json(results))

    if "bibtex" in formats:
        _safe_write("bibtex", "corrected.bib", lambda: render_bibtex(results))

    if "ris" in formats:
        _safe_write("ris", "corrected.ris", lambda: render_ris(results))

    if "html" in formats:
        _safe_write("html", "audit.html", lambda: render_html(
            results, style=style, source_name=source_name,
            crosscheck=crosscheck, duplicates=duplicates))

    if "docx" in formats:
        p = out_dir / "audit_tracked.docx"
        try:
            render_annotated_docx(results, style=style, out_path=p, source_name=source_name,
                                  crosscheck=crosscheck, duplicates=duplicates)
            written["docx"] = str(p)
        except Exception as e:
            import logging
            logging.getLogger("biblio-check.report").exception("docx render failed")
            written["docx_error"] = str(e)

    # Standalone clean bibliography is only useful when no other requested
    # format already contains one. Markdown, docx, bibtex, ris each carry the
    # clean bibliography natively (markdown/docx as a section; bibtex/ris are
    # themselves clean bibliographies). Only fall back to a separate .txt when
    # the user picked only JSON or no format that includes one.
    formats_with_clean = {"markdown", "docx", "bibtex", "ris"}
    if not (set(formats) & formats_with_clean):
        clean_path = out_dir / f"clean_bibliography.{style}.txt"
        clean_path.write_text(render_clean_bibliography(results, style), encoding="utf-8")
        written["clean_bibliography"] = str(clean_path)

    return written
